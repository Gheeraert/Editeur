from __future__ import annotations

import json
import sys
import unittest
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

from docx import Document as DocxDoc

from purh_editorial.config import load_settings
from purh_editorial.config.private_corpus import private_corpus_skip_reason, require_private_corpus
from purh_editorial.pipeline.step1 import Step1Options, Step1Pipeline

NS_WORD = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

# Au moins deux manuscrits réels doivent être présents dans manuscripts_raw/ pour que
# cette suite ait un sens (voir docs/CORPUS_ET_FIXTURES.md). Aucun nom de fichier réel
# n'apparaît dans ce module : les cas sont découverts dynamiquement dans le corpus privé
# local, jamais publiés dans ce dépôt.
MIN_EXPECTED_MANUSCRIPTS = 2


def _note_count(path: Path) -> int:
    with zipfile.ZipFile(path) as archive:
        try:
            root = ET.fromstring(archive.read("word/footnotes.xml"))
        except KeyError:
            return 0
    return sum(
        1
        for footnote in root.findall(".//w:footnote", NS_WORD)
        if footnote.attrib.get(f"{{{NS_WORD['w']}}}type")
        not in ("separator", "continuationSeparator")
    )


@unittest.skipIf(private_corpus_skip_reason() is not None, private_corpus_skip_reason() or "")
class PrivateEditorialCorpusTests(unittest.TestCase):
    """
    Tests exécutés uniquement quand PURH_PRIVATE_CORPUS_DIR pointe vers un corpus privé
    local installé (voir docs/CORPUS_ET_FIXTURES.md). Ignorés avec une raison explicite
    sur toute machine où cette variable n'est pas définie — en particulier en CI
    publique, qui ne doit jamais échouer pour cette seule raison.

    Ce module ne référence, ne contient et ne publie aucun contenu du corpus privé
    (ni titre, ni extrait, ni nom de fichier réel) : il découvre dynamiquement ce qui
    est présent localement.
    """

    @classmethod
    def setUpClass(cls) -> None:
        cls.private_paths = require_private_corpus()
        cls.manuscripts = sorted(cls.private_paths.manuscripts_raw.glob("*.docx"))
        if len(cls.manuscripts) < MIN_EXPECTED_MANUSCRIPTS:
            raise unittest.SkipTest(
                f"Corpus privé présent mais insuffisant : {len(cls.manuscripts)} "
                f"manuscrit(s) trouvé(s) sous manuscripts_raw/, {MIN_EXPECTED_MANUSCRIPTS} attendus."
            )

    def test_pipeline_runs_and_conserves_notes_on_each_manuscript(self) -> None:
        settings = load_settings()
        for source_path in self.manuscripts:
            with self.subTest(manuscript=source_path.name):
                expected_notes = _note_count(source_path)

                pipeline = Step1Pipeline(settings)
                options = Step1Options(decision_mode="deterministic")
                result = pipeline.run(source_path, options)

                self.assertEqual(
                    len(result.pipeline_result.source_document.notes),
                    expected_notes,
                )
                self.assertEqual(result.pipeline_result.report.warnings, [])

    def test_pipeline_is_idempotent_on_each_manuscript(self) -> None:
        settings = load_settings()
        runtime_dir = ROOT / "tests" / "_runtime"
        runtime_dir.mkdir(parents=True, exist_ok=True)

        for source_path in self.manuscripts:
            with self.subTest(manuscript=source_path.name):
                pass1_path = runtime_dir / f"private_idempotence_pass1_{source_path.name}"
                Step1Pipeline(settings).run(
                    source_path,
                    Step1Options(decision_mode="deterministic", output_path=pass1_path),
                )

                pass2_path = runtime_dir / f"private_idempotence_pass2_{source_path.name}"
                result2 = Step1Pipeline(settings).run(
                    pass1_path,
                    Step1Options(decision_mode="deterministic", output_path=pass2_path),
                )

                pass1_texts = [p.text for p in DocxDoc(str(pass1_path)).paragraphs]
                pass2_texts = [p.text for p in DocxDoc(str(pass2_path)).paragraphs]
                self.assertEqual(pass1_texts, pass2_texts)

                orthotypo_transformations = [
                    t
                    for t in result2.pipeline_result.report.transformations
                    if t.module == "orthotypo"
                ]
                self.assertEqual(orthotypo_transformations, [])

    def test_pipeline_matches_reference_corrected_copies(self) -> None:
        """
        Compare la sortie du pipeline à une copie corrigée de référence, si le corpus
        privé fournit un manifeste io_samples/pairs.json listant des couples
        {"raw": ..., "reference": ..., "check": "first_paragraph"}. Absent de ce dépôt
        public ; à fournir uniquement dans le corpus privé local.
        """
        manifest_path = self.private_paths.io_samples / "pairs.json"
        if not manifest_path.is_file():
            self.skipTest("Aucun manifeste io_samples/pairs.json dans le corpus privé local.")

        pairs = json.loads(manifest_path.read_text(encoding="utf-8"))
        settings = load_settings()
        for pair in pairs:
            with self.subTest(raw=pair["raw"]):
                raw_path = self.private_paths.io_samples / pair["raw"]
                reference_path = self.private_paths.io_samples / pair["reference"]

                pipeline = Step1Pipeline(settings)
                options = Step1Options(decision_mode="deterministic")
                result = pipeline.run(raw_path, options)
                produced = result.pipeline_result.source_document.blocks[0].text

                reference = DocxDoc(str(reference_path))
                expected = reference.paragraphs[0].text

                self.assertEqual(produced, expected)


if __name__ == "__main__":
    unittest.main()
