from __future__ import annotations

import hashlib
import os
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document as DocxDocument

from purh_editorial.services.word_review_service import WordReviewError, WordReviewService, _create_word_application
from purh_editorial.services.word_workspace_service import WordWorkspaceService


def _word_available() -> bool:
    if os.environ.get("PURH_RUN_WORD_INTEGRATION") != "1":
        raise unittest.SkipTest("PURH_RUN_WORD_INTEGRATION non activé")
    if sys.platform != "win32":
        raise unittest.SkipTest("Microsoft Word requiert Windows")
    try:
        app = _create_word_application()
    except WordReviewError as exc:
        raise AssertionError(
            f"PURH_RUN_WORD_INTEGRATION=1 mais Microsoft Word est indisponible : {exc}"
        ) from exc
    app.Quit()
    return True


@unittest.skipUnless(_word_available(), "Microsoft Word indisponible")
class WordWorkspaceIntegrationTests(unittest.TestCase):
    def test_real_workspace_opens_read_only_original_on_left(self) -> None:
        import pythoncom

        pythoncom.CoInitialize()
        with tempfile.TemporaryDirectory() as root:
            original = Path(root) / "original.docx"
            candidate = Path(root) / "candidate.docx"
            review = Path(root) / "review.docx"
            original_doc = DocxDocument(); original_doc.add_paragraph("Le XVIIème siècle."); original_doc.save(original)
            candidate_doc = DocxDocument(); candidate_doc.add_paragraph("Le XVIIe siècle."); candidate_doc.save(candidate)
            original_sha = hashlib.sha256(original.read_bytes()).hexdigest()
            WordReviewService().create_review_document(
                original_path=original, revised_path=candidate, output_path=review
            )
            session = WordWorkspaceService().open_workspace(original_path=original, review_path=review)
            try:
                self.assertTrue(bool(session.original_document.ReadOnly))
                self.assertFalse(bool(session.review_document.ReadOnly))
                self.assertLess(session.original_window.Left, session.review_window.Left)
                self.assertTrue(session.state.original_on_left)
                self.assertTrue(
                    session.state.synchronized_scrolling,
                    {
                        "warnings": session.state.warnings,
                        "strategy": session.state.side_by_side_strategy,
                        "attempts": [
                            {
                                "strategy": item.strategy,
                                "error_type": item.error_type,
                                "hresult": item.hresult,
                                "error_message": item.error_message,
                            }
                            for item in session.state.side_by_side_attempts
                        ],
                    },
                )
                self.assertEqual(hashlib.sha256(original.read_bytes()).hexdigest(), original_sha)
            finally:
                session.close_owned_documents(save_review=False)
                session.quit_application()
        pythoncom.CoUninitialize()


if __name__ == "__main__":
    unittest.main()
