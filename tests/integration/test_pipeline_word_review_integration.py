from __future__ import annotations

import hashlib
import os
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

from purh_editorial.config import load_settings
from purh_editorial.pipeline.step1 import Step1Options, Step1Pipeline
from purh_editorial.services.word_review_service import (
    WordReviewError,
    _create_word_application,
    document_contains_tracked_changes,
)


def _word_is_available() -> bool:
    if os.environ.get("PURH_RUN_WORD_INTEGRATION") != "1" or sys.platform != "win32":
        return False
    try:
        application = _create_word_application()
    except WordReviewError:
        return False
    try:
        return True
    finally:
        application.Quit()


@unittest.skipUnless(
    os.environ.get("PURH_RUN_WORD_INTEGRATION") == "1",
    "test Word reel ignore sans PURH_RUN_WORD_INTEGRATION=1",
)
@unittest.skipUnless(sys.platform == "win32", "Microsoft Word COM requiert Windows")
@unittest.skipUnless(_word_is_available(), "Microsoft Word indisponible")
class PipelineWordReviewIntegrationTests(unittest.TestCase):
    def test_pipeline_exports_candidate_then_native_word_review(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp = Path(tmpdir)
            original = tmp / "original.docx"
            candidate = tmp / "candidate.docx"
            review = tmp / "revision.docx"

            source_document = Document()
            source_document.add_paragraph("Le XVII\u00e8me si\u00e8cle.")
            source_document.save(original)
            original_sha256 = hashlib.sha256(original.read_bytes()).hexdigest()

            result = Step1Pipeline(settings=load_settings()).run(
                original,
                Step1Options(output_path=candidate, word_review_output_path=review),
            )

            self.assertEqual(hashlib.sha256(original.read_bytes()).hexdigest(), original_sha256)
            self.assertTrue(candidate.exists())
            self.assertIn("Le xviie si\u00e8cle.", "\n".join(p.text for p in Document(candidate).paragraphs))
            self.assertTrue(review.exists())
            self.assertTrue(document_contains_tracked_changes(review))
            self.assertIsNotNone(result.word_review_result)
            self.assertTrue(result.word_review_result.has_tracked_changes)
            module_run = next(
                run for run in result.pipeline_result.report.module_runs if run.module_name == "word_review"
            )
            self.assertEqual(module_run.status, "success")


if __name__ == "__main__":
    unittest.main()
