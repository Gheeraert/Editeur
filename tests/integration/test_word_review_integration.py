from __future__ import annotations

import os
import sys
import tempfile
import unittest
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

from purh_editorial.services.word_review_service import (
    WordReviewError,
    WordReviewService,
    document_contains_tracked_changes,
)


@unittest.skipUnless(
    os.environ.get("PURH_RUN_WORD_INTEGRATION") == "1",
    "test Word reel ignore sans PURH_RUN_WORD_INTEGRATION=1",
)
@unittest.skipUnless(sys.platform == "win32", "Microsoft Word COM requiert Windows")
class WordReviewIntegrationTests(unittest.TestCase):
    def test_real_word_comparison_creates_tracked_changes(self) -> None:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover
            self.skipTest(f"python-docx indisponible : {exc}")

        with tempfile.TemporaryDirectory() as tmpdir:
            tmp = Path(tmpdir)
            original = tmp / "original.docx"
            candidate = tmp / "candidate.docx"
            output = tmp / "revision.docx"

            original_doc = Document()
            original_doc.add_paragraph("Il n'y a pas de probleme...")
            original_doc.add_paragraph("Le XVIIeme siecle.")
            original_doc.save(original)

            candidate_doc = Document()
            candidate_doc.add_paragraph("Il n\u2019y a pas de probleme\u2026")
            candidate_doc.add_paragraph("Le XVIIe\u00a0siecle.")
            candidate_doc.save(candidate)

            try:
                result = WordReviewService().create_review_document(
                    original_path=original,
                    revised_path=candidate,
                    output_path=output,
                )
            except WordReviewError as exc:
                self.skipTest(f"Microsoft Word indisponible pour ce test : {exc}")

            self.assertTrue(result.output_path.exists())
            self.assertTrue(document_contains_tracked_changes(output))


if __name__ == "__main__":
    unittest.main()
