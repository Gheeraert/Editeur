from __future__ import annotations

import hashlib
import os
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_COLOR_INDEX

from purh_editorial.model import Diagnostic, Document, Evidence, Paragraph, ProcessingReport
from purh_editorial.services.word_review_annotation_service import (
    WordReviewAnnotationService,
    document_contains_highlights,
    document_contains_word_comments,
)
from purh_editorial.services.word_review_service import WordReviewError, WordReviewService, _create_word_application, document_contains_tracked_changes


def _word_available() -> bool:
    if os.environ.get("PURH_RUN_WORD_INTEGRATION") != "1": raise unittest.SkipTest("PURH_RUN_WORD_INTEGRATION non activé")
    if sys.platform != "win32": raise unittest.SkipTest("Microsoft Word requiert Windows")
    try: app = _create_word_application()
    except WordReviewError as exc: raise AssertionError(f"PURH_RUN_WORD_INTEGRATION=1 mais Microsoft Word est indisponible : {exc}") from exc
    app.Quit(); return True


@unittest.skipUnless(_word_available(), "Microsoft Word indisponible ou intégration désactivée")
class WordReviewAnnotationIntegrationTests(unittest.TestCase):
    def test_real_annotation_preserves_revisions_and_highlights(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir); original = root / "original.docx"; candidate = root / "candidate.docx"; review = root / "review.docx"
            doc = DocxDocument(); doc.add_paragraph("Le XVIIème siècle."); doc.add_paragraph("Ce passage doit être vérifié."); doc.save(original)
            candidate_doc = DocxDocument(); candidate_doc.add_paragraph("Le XVIIe siècle."); paragraph = candidate_doc.add_paragraph(); run = paragraph.add_run("Ce passage doit être vérifié."); run.font.highlight_color = WD_COLOR_INDEX.YELLOW; candidate_doc.save(candidate)
            original_sha = hashlib.sha256(original.read_bytes()).hexdigest(); candidate_sha = hashlib.sha256(candidate.read_bytes()).hexdigest()
            WordReviewService().create_review_document(original_path=original, revised_path=candidate, output_path=review)
            pivot = Document(document_id="d", source_path=str(original), source_format="docx", blocks=[Paragraph(block_id="p1", text="Ce passage doit être vérifié.")])
            report = ProcessingReport(report_id="r", document_id="d", diagnostics=[Diagnostic("d1", "review", "warning", "review", "Vérifier ce passage.", "p1", Evidence(excerpt="Ce passage doit être vérifié."))])
            result = WordReviewAnnotationService().annotate_review_document(review_path=review, document=pivot, report=report)
            self.assertEqual(result.comments_requested, 1); self.assertEqual(result.comments_added, 1); self.assertEqual(result.comments_skipped, 0)
            self.assertEqual(hashlib.sha256(original.read_bytes()).hexdigest(), original_sha)
            self.assertEqual(hashlib.sha256(candidate.read_bytes()).hexdigest(), candidate_sha)
            self.assertTrue(document_contains_tracked_changes(review)); self.assertTrue(document_contains_highlights(review)); self.assertTrue(document_contains_word_comments(review))
            with zipfile.ZipFile(review) as archive: self.assertIn("Vérifier ce passage.".encode(), archive.read("word/comments.xml"))
            app = _create_word_application()
            opened = None
            try:
                opened = app.Documents.Open(
                    FileName=str(review), ReadOnly=True, AddToRecentFiles=False, Visible=False
                )
            finally:
                if opened is not None:
                    opened.Close(SaveChanges=False)
                app.Quit()


if __name__ == "__main__": unittest.main()
