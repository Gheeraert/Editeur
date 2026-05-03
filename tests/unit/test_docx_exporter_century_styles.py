from __future__ import annotations

import sys
import unittest
import uuid
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document as DocxDoc

ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

from purh_editorial.io.docx_exporter import DocxExporter
from purh_editorial.io.importer_registry import ImporterRegistry
from purh_editorial.model import Document, InlineSpan, InlineStyle, LineatedBlock, Paragraph
from purh_editorial.services.orthotypo_service import OrthotypoService
from tests.helpers.docx_factory import create_table_docx

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _q(local_name: str) -> str:
    return f"{{{W_NS}}}{local_name}"


class DocxExporterCenturyStylesTests(unittest.TestCase):
    @staticmethod
    def _runtime_path(filename: str) -> Path:
        runtime_dir = ROOT / "tests" / "_runtime"
        runtime_dir.mkdir(parents=True, exist_ok=True)
        return runtime_dir / filename

    def _minimal_template_docx(self) -> Path:
        template_path = self._runtime_path(f"template_{uuid.uuid4().hex}.docx")
        doc = DocxDoc()
        doc.save(template_path)
        return template_path

    def _export_and_read_document_xml(self, document: Document) -> ET.Element:
        template = self._minimal_template_docx()
        output = self._runtime_path(f"export_{uuid.uuid4().hex}.docx")
        DocxExporter(template_path=template).export(document, output)
        with zipfile.ZipFile(output, "r") as archive:
            xml_bytes = archive.read("word/document.xml")
        return ET.fromstring(xml_bytes)

    @staticmethod
    def _runs_with_text(root: ET.Element) -> list[tuple[str, ET.Element]]:
        runs: list[tuple[str, ET.Element]] = []
        for run in root.findall(f".//{_q('r')}"):
            texts = [t.text or "" for t in run.findall(_q("t"))]
            run_text = "".join(texts)
            if run_text:
                runs.append((run_text, run))
        return runs

    def test_century_is_exported_with_small_caps_and_superscript(self) -> None:
        source = Document(
            document_id="doc1",
            source_path="tests/fixtures/minimal_source.txt",
            source_format="txt",
            blocks=[Paragraph(block_id="p1", text="XIXe siècle")],
        )
        styled, _ = OrthotypoService().apply(source)
        root = self._export_and_read_document_xml(styled)
        runs = self._runs_with_text(root)

        roman_run = next((run for text, run in runs if text == "xix"), None)
        self.assertIsNotNone(roman_run)
        self.assertIsNotNone(roman_run.find(f"./{_q('rPr')}/{_q('smallCaps')}"))

        exponent_run = next((run for text, run in runs if text == "e"), None)
        self.assertIsNotNone(exponent_run)
        vert_align = exponent_run.find(f"./{_q('rPr')}/{_q('vertAlign')}")
        self.assertIsNotNone(vert_align)
        self.assertEqual(vert_align.attrib.get(_q("val")), "superscript")

        siecle_run = next((run for text, run in runs if "siècle" in text), None)
        self.assertIsNotNone(siecle_run)
        self.assertIsNone(siecle_run.find(f"./{_q('rPr')}/{_q('smallCaps')}"))
        self.assertIsNone(siecle_run.find(f"./{_q('rPr')}/{_q('vertAlign')}"))

    def test_century_keeps_italic_in_exported_runs(self) -> None:
        source = Document(
            document_id="doc2",
            source_path="tests/fixtures/minimal_source.txt",
            source_format="txt",
            blocks=[
                Paragraph(
                    block_id="p1",
                    text="XVIIe siècle",
                    inlines=[InlineSpan(text="XVIIe siècle", style=InlineStyle(italic=True))],
                )
            ],
        )
        styled, _ = OrthotypoService().apply(source)
        root = self._export_and_read_document_xml(styled)
        runs = self._runs_with_text(root)

        roman_run = next((run for text, run in runs if text == "xvii"), None)
        self.assertIsNotNone(roman_run.find(f"./{_q('rPr')}/{_q('i')}"))
        self.assertIsNotNone(roman_run.find(f"./{_q('rPr')}/{_q('smallCaps')}"))

        exponent_run = next((run for text, run in runs if text == "e"), None)
        self.assertIsNotNone(exponent_run.find(f"./{_q('rPr')}/{_q('i')}"))
        self.assertIsNotNone(exponent_run.find(f"./{_q('rPr')}/{_q('vertAlign')}"))

    def test_flat_block_gets_styled_and_exported(self) -> None:
        source = Document(
            document_id="doc3",
            source_path="tests/fixtures/minimal_source.txt",
            source_format="txt",
            blocks=[Paragraph(block_id="p1", text="XIXe siècle")],
        )
        styled, _ = OrthotypoService().apply(source)
        self.assertTrue(styled.blocks[0].inlines)

        root = self._export_and_read_document_xml(styled)
        runs = self._runs_with_text(root)
        self.assertTrue(any(text == "xix" for text, _ in runs))
        self.assertTrue(any(text == "e" for text, _ in runs))

    def test_export_preserves_table_and_body_order(self) -> None:
        runtime_input = self._runtime_path(f"table_{uuid.uuid4().hex}.docx")
        create_table_docx(runtime_input)
        imported = ImporterRegistry().load_document(runtime_input)

        template = self._minimal_template_docx()
        output = self._runtime_path(f"export_table_{uuid.uuid4().hex}.docx")
        DocxExporter(template_path=template).export(imported, output)

        with zipfile.ZipFile(output, "r") as archive:
            xml_bytes = archive.read("word/document.xml")
        root = ET.fromstring(xml_bytes)
        body = root.find(_q("body"))
        self.assertIsNotNone(body)
        body_tags = [node.tag for node in list(body) if node.tag != _q("sectPr")]
        self.assertEqual(body_tags, [_q("p"), _q("tbl"), _q("p")])

        tbl_nodes = body.findall(_q("tbl"))
        self.assertEqual(len(tbl_nodes), 1)
        first_cell_text = "".join(t.text or "" for t in tbl_nodes[0].findall(f".//{_q('tc')}//{_q('t')}"))
        self.assertIn("VIII,Pr.,", first_cell_text)
        self.assertIn("VIII;Pr.,18", first_cell_text)

    def test_pipeline_structure_does_not_promote_table_content_to_heading(self) -> None:
        runtime_input = self._runtime_path(f"table_no_heading_{uuid.uuid4().hex}.docx")
        create_table_docx(runtime_input)
        imported = ImporterRegistry().load_document(runtime_input)
        heading_count = sum(1 for block in imported.blocks if block.block_type == "heading")
        self.assertEqual(heading_count, 0)

    def test_lineated_block_inlines_preserve_word_line_break_and_italic(self) -> None:
        document = Document(
            document_id="doc-lineated-inline-break",
            source_path="tests/fixtures/minimal_source.txt",
            source_format="txt",
            blocks=[
                LineatedBlock(
                    block_id="lb1",
                    text="Je vois Phèdre venir.\nDeuxième vers.",
                    inlines=[
                        InlineSpan(text="Je vois "),
                        InlineSpan(text="Phèdre", style=InlineStyle(italic=True)),
                        InlineSpan(text=" venir."),
                        InlineSpan(text="", kind="line_break"),
                        InlineSpan(text="Deuxième vers."),
                    ],
                    attributes={
                        "semantic": {
                            "role": "lineated_block",
                            "layout_kind": "lineated_block",
                            "lineation": "lineated",
                            "lines": ["Je vois Phèdre venir.", "Deuxième vers."],
                        }
                    },
                )
            ],
        )
        root = self._export_and_read_document_xml(document)
        br_nodes = root.findall(f".//{_q('br')}")
        self.assertTrue(br_nodes)

        runs = self._runs_with_text(root)
        phedre_run = next((run for text, run in runs if "Phèdre" in text), None)
        self.assertIsNotNone(phedre_run)
        italic = phedre_run.find(f"./{_q('rPr')}/{_q('i')}")
        self.assertIsNotNone(italic)


if __name__ == "__main__":
    unittest.main()
