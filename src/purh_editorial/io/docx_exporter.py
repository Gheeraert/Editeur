from __future__ import annotations

import io
import re
import zipfile
from pathlib import Path
from xml.sax.saxutils import escape as xml_escape
from xml.etree import ElementTree as ET

from docx import Document as DocxDoc
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt

from purh_editorial.config.private_corpus import resolve_private_corpus_dir
from purh_editorial.model import Document, ImageAsset, ImageOccurrence, InlineSpan, Note
from purh_editorial.model.semantics import extract_verse_lines, is_canonical_lineated_block

# ── Gabarit Métopes ───────────────────────────────────────────────────────────
# Le gabarit réel (Commons-publishing-Metopes.dotm) est un document interne PURH,
# jamais suivi dans le dépôt public (voir docs/CORPUS_ET_FIXTURES.md). Résolution,
# par ordre de priorité :
#   1. corpus privé local (PURH_PRIVATE_CORPUS_DIR), gabarit réel avec macros/styles ;
#   2. ancien chemin relatif au dépôt (compatibilité, généralement absent en public) ;
#   3. gabarit public minimal (fixtures/templates/), styles Word par défaut
#      uniquement — suffisant pour les tests publics, pas pour une édition réelle.
# Une erreur claire est levée si aucun des trois n'est trouvé.
_TEMPLATE_RELATIVE_PATH = "sources/editorial_rules/metopes_template_word/Commons-publishing-Metopes.dotm"
_REPO_ROOT = Path(__file__).parent.parent.parent.parent
_PUBLIC_TEMPLATE_PATH = _REPO_ROOT / "fixtures" / "templates" / "metopes_template_public.docx"


def resolve_default_template(explicit_path: Path | None = None) -> Path:
    """Résout le gabarit Métopes à utiliser (voir ordre de priorité ci-dessus).

    Lève FileNotFoundError avec un message explicite si aucun gabarit n'est
    disponible, plutôt que de laisser une exception opaque remonter plus tard."""
    if explicit_path is not None:
        if explicit_path.is_file():
            return explicit_path
        raise FileNotFoundError(f"Gabarit Métopes explicite introuvable : {explicit_path}")

    private_root = resolve_private_corpus_dir()
    if private_root is not None:
        candidate = private_root / _TEMPLATE_RELATIVE_PATH
        if candidate.is_file():
            return candidate

    legacy_candidate = _REPO_ROOT / _TEMPLATE_RELATIVE_PATH
    if legacy_candidate.is_file():
        return legacy_candidate

    if _PUBLIC_TEMPLATE_PATH.is_file():
        return _PUBLIC_TEMPLATE_PATH

    raise FileNotFoundError(
        "Aucun gabarit Métopes disponible : configurez PURH_PRIVATE_CORPUS_DIR "
        "avec le gabarit réel, ou fournissez template_path explicitement. "
        f"Le gabarit public minimal ({_PUBLIC_TEMPLATE_PATH}) est aussi absent."
    )


def _resolve_default_template() -> Path:
    return resolve_default_template()



# Content-type .dotm → .docx (python-docx refuse les templates macro)
_CT_DOTM = b"application/vnd.ms-word.template.macroEnabledTemplate.main+xml"
_CT_DOCX = b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"

# Fichiers VBA à exclure lors de la conversion .dotm -> .docx.
_VBA_FILES = frozenset({"word/vbaProject.bin", "word/vbaData.xml", "word/_rels/vbaProject.bin.rels"})
# Le ruban Métopes appelle des callbacks VBA : il faut le retirer avec les macros.
_CUSTOM_UI_PREFIXES = ("customUI/",)

# Retire les entrées VBA dans [Content_Types].xml.
_VBA_CT_RE = re.compile(rb'<Override[^>]+/word/vba[^>]+/>')
_VBA_DEFAULT_CT_RE = re.compile(
    rb'<Default\b(?=[^>]*\bExtension="bin")'
    rb'(?=[^>]*\bContentType="application/vnd\.ms-office\.vbaProject")[^>]*/>'
)
# Retire les relations vers le projet VBA ou vers le ruban customUI.
_VBA_REL_RE = re.compile(rb'<Relationship[^>]+vbaProject[^>]*/>')
_CUSTOM_UI_REL_RE = re.compile(rb'<Relationship[^>]+(?:customUI|ui/extensibility)[^>]*/>')

# Namespaces OOXML
W  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML = "http://www.w3.org/XML/1998/namespace"

# ── Correspondance couleurs de surlignage ─────────────────────────────────────
_HIGHLIGHT_MAP: dict[str, WD_COLOR_INDEX] = {
    "orthotypo":             WD_COLOR_INDEX.YELLOW,
    "footnote":              WD_COLOR_INDEX.BRIGHT_GREEN,
    "biblio":                WD_COLOR_INDEX.TURQUOISE,
    "ai":                    WD_COLOR_INDEX.PINK,
    "ai_structure":          WD_COLOR_INDEX.VIOLET,
    "exploratory_structure": WD_COLOR_INDEX.DARK_YELLOW,
    # Modes non-exploratoires : traité automatiquement (transform zone)
    "structure_applied":     WD_COLOR_INDEX.TEAL,
    # Tous modes : suspect non traité, signalement manuel requis
    "suspect_unhandled":     WD_COLOR_INDEX.RED,
}

# ── Polices PURH (substituts des polices InDesign) ───────────────────────────
# Corps du texte → Garamond (empattement, analogue Chaparral Pro)
# Titraille      → Calibri  (sans empattement, analogue Josefin Sans)
_FONT_BODY = "Garamond"
_FONT_HEAD = "Calibri"
_HEADING_FONT_SIZES_PT: dict[int, int] = {1: 16, 2: 14, 3: 13}
_QUOTE_FONT_SIZE_PT = 11
_QUOTE_LEFT_INDENT_CM = 0.7

# Styles de corps (reçoivent Garamond)
_BODY_STYLE_NAMES: frozenset[str] = frozenset({
    "Normal", "footnote text", "endnote text",
    "TEI_quote", "TEI_quote2", "TEI_quote_nested", "TEI_quote_continuation",
    "TEI_bibl_reference", "TEI_epigraph", "TEI_acknowledgment",
    "TEI_dedication", "TEI_paragraph_lead", "TEI_paragraph_consecutive",
    "TEI_abstract", "TEI_keywords", "TEI_verse", "TEI_figure_caption",
    "TEI_figure_credits", "TEI_figure_alternative", "TEI_aut:",
    "TEI_note:", "TEI_localpara",
})

# Styles de titre (reçoivent Calibri)
_HEAD_STYLE_NAMES: frozenset[str] = frozenset({
    "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5",
    "heading 1", "heading 2", "heading 3", "heading 4", "heading 5",
    "TEI_bibl_start", "TEI_appendix_start",
})
_FIRST_LINE_INDENT_CM = 0.5


# ── Conversion .dotm → flux .docx ────────────────────────────────────────────

def _dotm_to_docx_bytes(dotm_path: Path) -> io.BytesIO:
    """Retourne le contenu du .dotm converti en .docx sans macros ni ruban VBA."""
    buf = io.BytesIO()
    with zipfile.ZipFile(dotm_path, "r") as src, \
         zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            if item.filename in _VBA_FILES or item.filename.startswith(_CUSTOM_UI_PREFIXES):
                continue
            data = src.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = data.replace(_CT_DOTM, _CT_DOCX)
                data = _VBA_CT_RE.sub(b"", data)
                data = _VBA_DEFAULT_CT_RE.sub(b"", data)
            elif item.filename.endswith(".rels"):
                data = _VBA_REL_RE.sub(b"", data)
                data = _CUSTOM_UI_REL_RE.sub(b"", data)
            dst.writestr(item, data)
    buf.seek(0)
    return buf


# ── Helpers python-docx ───────────────────────────────────────────────────────

def _add_line_break(paragraph) -> None:
    """Insère un saut de ligne manuel (Shift+Entrée) dans un paragraphe Word."""
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    r.append(br)
    paragraph._p.append(r)


def _add_page_break(doc: DocxDoc) -> None:
    """Insère un saut de page (paragraphe vide avec <w:br w:type='page'>)."""
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    _insert_body_element(doc, p)


def _add_footnote_reference(paragraph, footnote_id: int) -> None:
    """Insère un appel de note (renvoi superscript) dans un paragraphe."""
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "Appelnotedebasdep")
    rpr.append(rstyle)
    r.append(rpr)
    ref = OxmlElement("w:footnoteReference")
    ref.set(qn("w:id"), str(footnote_id))
    r.append(ref)
    paragraph._p.append(r)


def _add_run_with_style(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    small_caps: bool = False,
    superscript: bool = False,
    subscript: bool = False,
    highlight: WD_COLOR_INDEX | None = None,
    font_name: str | None = None,
    font_size_pt: int | None = None,
) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if small_caps:
        run.font.small_caps = True
    if superscript:
        run.font.superscript = True
    if subscript:
        run.font.subscript = True
    if highlight is not None:
        run.font.highlight_color = highlight
    if font_name:
        run.font.name = font_name
    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)


def _inline_highlight(span: InlineSpan) -> WD_COLOR_INDEX | None:
    key = span.attributes.get("highlight_color")
    return _HIGHLIGHT_MAP.get(key) if key else None


def _add_image_run(para, occurrence: ImageOccurrence, image_assets: dict[str, ImageAsset]) -> None:
    """Réinsère une image incorporée (image externe non incorporée : rien à
    réinjecter, déjà diagnostiquée à l'import). Les images ancrées sont
    normalisées en inline (voir ImageOccurrence.anchor_normalized_to_inline)."""
    asset = image_assets.get(occurrence.asset_ref) if occurrence.asset_ref else None
    if asset is None:
        return
    run = para.add_run()
    kwargs: dict = {}
    if occurrence.width_emu:
        kwargs["width"] = Emu(occurrence.width_emu)
    if occurrence.height_emu:
        kwargs["height"] = Emu(occurrence.height_emu)
    inline_shape = run.add_picture(io.BytesIO(asset.data), **kwargs)
    if occurrence.alt_text:
        inline_shape._inline.docPr.set("descr", occurrence.alt_text)
    if occurrence.title:
        inline_shape._inline.docPr.set("title", occurrence.title)


def _add_paragraph(
    doc: DocxDoc,
    block,
    note_id_map: dict[str, int],
    image_occurrences_by_id: dict[str, ImageOccurrence],
    image_assets: dict[str, ImageAsset],
) -> None:
    """Ajoute un paragraphe au document Word avec son style Métopes."""
    style_name = block.attributes.get("metopes_style", "Normal")
    try:
        para = doc.add_paragraph(style=style_name)
    except (KeyError, ValueError):
        para = doc.add_paragraph(style="Normal")

    if style_name == "Normal":
        para.paragraph_format.first_line_indent = Cm(_FIRST_LINE_INDENT_CM)

    # Retrait suspendu pour les entrées bibliographiques (≈ 5 mm, charte PURH)
    if style_name == "TEI_bibl_reference":
        para.paragraph_format.left_indent = Cm(0.5)
        para.paragraph_format.first_line_indent = Cm(-0.5)

    heading_level = 0
    if block.block_type == "heading":
        raw_level = block.attributes.get("heading_level", 0)
        try:
            heading_level = int(raw_level or 0)
        except (TypeError, ValueError):
            heading_level = 1
    heading_size = _HEADING_FONT_SIZES_PT.get(min(max(heading_level, 1), 3)) if heading_level else None
    is_lineated_block = is_canonical_lineated_block(block)
    is_quote_block = block.block_type == "quote_block"
    if is_quote_block or block.block_type == "lineated_block":
        para.paragraph_format.left_indent = Cm(_QUOTE_LEFT_INDENT_CM)
        para.paragraph_format.first_line_indent = None
        para.paragraph_format.line_spacing = 1.0
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT if is_lineated_block else WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    if block.inlines:
        block_hl = _HIGHLIGHT_MAP.get(block.attributes.get("highlight_color", ""), None)
        for span in block.inlines:
            if span.kind == "line_break":
                _add_line_break(para)
                continue
            if span.kind == "note_call" and span.note_ref:
                fn_id = note_id_map.get(span.note_ref)
                if fn_id is not None:
                    _add_footnote_reference(para, fn_id)
            elif span.kind == "image":
                occ = image_occurrences_by_id.get(span.attributes.get("occurrence_id", ""))
                if occ is not None:
                    _add_image_run(para, occ, image_assets)
            elif span.kind == "complex_object":
                # Objet non reconstruit a l'export (graphique/SmartArt/OLE/equation) ;
                # deja detecte et diagnostique a l'import, jamais confondu avec une image.
                continue
            else:
                span_hl = _inline_highlight(span)
                _add_run_with_style(
                    para,
                    span.text,
                    bold=span.style.bold,
                    italic=span.style.italic,
                    small_caps=span.style.small_caps,
                    superscript=span.style.superscript,
                    subscript=span.style.subscript,
                    highlight=span_hl if span_hl is not None else block_hl,
                    font_size_pt=_QUOTE_FONT_SIZE_PT if (is_quote_block or is_lineated_block) else heading_size,
                )
    else:
        hl = _HIGHLIGHT_MAP.get(block.attributes.get("highlight_color", ""), None)
        if is_lineated_block:
            lines = extract_verse_lines(block)
            for i, line in enumerate(lines):
                if i > 0:
                    _add_line_break(para)
                _add_run_with_style(
                    para,
                    line.strip(),
                    highlight=hl,
                    font_size_pt=_QUOTE_FONT_SIZE_PT,
                )
        else:
            _add_run_with_style(
                para,
                block.text,
                highlight=hl,
                font_size_pt=_QUOTE_FONT_SIZE_PT if is_quote_block else heading_size,
            )


def _insert_body_element(doc: DocxDoc, element) -> None:
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is None:
        body.append(element)
        return
    body.insert(body.index(sect_pr), element)


def _add_raw_table(doc: DocxDoc, block, image_assets: dict[str, ImageAsset]) -> None:
    raw_ooxml = str(block.attributes.get("table_ooxml", "") or "").strip()
    if not raw_ooxml:
        return
    for ref in block.attributes.get("table_image_refs") or []:
        old_rid = ref.get("rid")
        asset_id = ref.get("asset_id")
        if not old_rid or not asset_id:
            continue
        asset = image_assets.get(asset_id)
        if asset is None:
            continue
        new_rid, _image = doc.part.get_or_add_image(io.BytesIO(asset.data))
        # Le préfixe de la relation (r:, ns4:, ...) dépend de la sérialisation
        # ElementTree au moment de l'import (voir docx_importer.py) et n'est pas
        # stable : on remplace par valeur exacte de rId, quel que soit le préfixe.
        raw_ooxml = re.sub(
            rf'((?:\w+:)?(?:embed|link|id)=")({re.escape(old_rid)})(")',
            rf'\g<1>{new_rid}\g<3>',
            raw_ooxml,
        )
    table_element = parse_xml(raw_ooxml.encode("utf-8"))
    _insert_body_element(doc, table_element)


# ── Injection des notes de bas de page (post-sauvegarde) ─────────────────────

_HL_TO_WVAL = {
    "orthotypo":             "yellow",
    "footnote":              "green",
    "biblio":                "cyan",
    "ai":                    "magenta",
    "exploratory_structure": "darkYellow",
    "structure_applied":     "darkCyan",
    "suspect_unhandled":     "red",
}

_FOOTNOTES_XML_DECL = b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\r\n'
_FOOTNOTES_FALLBACK_ROOT = (
    '<w:footnotes '
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
    'xmlns:xml="http://www.w3.org/XML/1998/namespace">'
    '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
    '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
    '</w:footnotes>'
).encode("utf-8")


def _xml_text(value: str) -> str:
    return xml_escape(value, {"\r": "&#13;"})


def _xml_attr(value: object) -> str:
    return xml_escape(str(value), {'"': '&quot;', "'": '&apos;'})


def _note_rpr_xml(
    *,
    superscript: bool = False,
    bold: bool = False,
    italic: bool = False,
    small_caps: bool = False,
    highlight_key: str | None = None,
) -> str:
    parts = [
        '<w:rPr>',
        '<w:rStyle w:val="NotedebasdepageCar"/>',
        f'<w:rFonts w:ascii="{_xml_attr(_FONT_BODY)}" w:hAnsi="{_xml_attr(_FONT_BODY)}"/>',
    ]
    if bold:
        parts.append('<w:b/>')
    if italic:
        parts.append('<w:i/>')
    if small_caps:
        parts.append('<w:smallCaps w:val="true"/>')
    if superscript:
        parts.append('<w:vertAlign w:val="superscript"/>')
    hl_val = _HL_TO_WVAL.get(highlight_key or "")
    if hl_val:
        parts.append(f'<w:highlight w:val="{_xml_attr(hl_val)}"/>')
    parts.append('</w:rPr>')
    return "".join(parts)


def _note_text_run_xml(
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    small_caps: bool = False,
    superscript: bool = False,
    highlight_key: str | None = None,
) -> str:
    if not text:
        return ""
    return (
        '<w:r>'
        + _note_rpr_xml(
            bold=bold,
            italic=italic,
            small_caps=small_caps,
            superscript=superscript,
            highlight_key=highlight_key,
        )
        + f'<w:t xml:space="preserve">{_xml_text(text)}</w:t>'
        + '</w:r>'
    )


# ── Images incorporées dans les notes (injection post-sauvegarde) ────────────

_CONTENT_TYPE_EXT = {
    "image/png": "png", "image/jpeg": "jpg", "image/gif": "gif", "image/bmp": "bmp",
    "image/tiff": "tif", "image/svg+xml": "svg", "image/x-emf": "emf", "image/x-wmf": "wmf",
}
_DEFAULT_IMAGE_EXTENT_EMU = 914400  # 1 pouce, valeur de repli si aucune dimension connue.


def _note_image_drawing_xml(rid: str, occurrence: ImageOccurrence, doc_pr_id: int) -> str:
    cx = occurrence.width_emu or _DEFAULT_IMAGE_EXTENT_EMU
    cy = occurrence.height_emu or _DEFAULT_IMAGE_EXTENT_EMU
    name = _xml_attr(occurrence.title or "Image")
    descr_attr = f' descr="{_xml_attr(occurrence.alt_text)}"' if occurrence.alt_text else ""
    title_attr = f' title="{_xml_attr(occurrence.title)}"' if occurrence.title else ""
    return (
        '<w:r><w:drawing>'
        '<wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        f'<wp:docPr id="{doc_pr_id}" name="{name}"{descr_attr}{title_attr}/>'
        '<wp:cNvGraphicFramePr>'
        '<a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/>'
        '</wp:cNvGraphicFramePr>'
        '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:nvPicPr>'
        f'<pic:cNvPr id="{doc_pr_id}" name="{name}"/>'
        '<pic:cNvPicPr/>'
        '</pic:nvPicPr>'
        '<pic:blipFill>'
        f'<a:blip xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" r:embed="{rid}"/>'
        '<a:stretch><a:fillRect/></a:stretch>'
        '</pic:blipFill>'
        '<pic:spPr>'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        '</pic:spPr>'
        '</pic:pic>'
        '</a:graphicData>'
        '</a:graphic>'
        '</wp:inline>'
        '</w:drawing></w:r>'
    )


def _build_note_image_plan(
    notes: list[Note],
    image_occurrences_by_id: dict[str, ImageOccurrence],
    image_assets: dict[str, ImageAsset],
    existing_rids: set[str],
) -> dict[str, dict]:
    """Attribue un rId et un nom de fichier media a chaque image de note a inserer.

    Les octets identiques (meme asset) partagent le meme fichier media, mais chaque
    occurrence garde sa propre relation (les relations sont locales a footnotes.xml,
    rien n'empeche plusieurs rId de pointer vers la meme cible)."""
    plan: dict[str, dict] = {}
    used = set(existing_rids)
    next_num = 1
    media_filename_by_asset: dict[str, str] = {}

    def _fresh_rid() -> str:
        nonlocal next_num
        while f"rId{next_num}" in used:
            next_num += 1
        rid = f"rId{next_num}"
        used.add(rid)
        return rid

    for note in notes:
        for span in note.inlines:
            if span.kind != "image":
                continue
            occ = image_occurrences_by_id.get(span.attributes.get("occurrence_id", ""))
            if occ is None or not occ.asset_ref:
                continue
            asset = image_assets.get(occ.asset_ref)
            if asset is None:
                continue
            filename = media_filename_by_asset.get(asset.asset_id)
            if filename is None:
                ext = _CONTENT_TYPE_EXT.get(asset.content_type, "png")
                filename = f"media/note_{asset.asset_id}.{ext}"
                media_filename_by_asset[asset.asset_id] = filename
            plan[occ.occurrence_id] = {"rid": _fresh_rid(), "filename": filename, "asset": asset, "occurrence": occ}
    return plan


def _build_footnote_xml(
    note: Note,
    footnote_id: int,
    note_image_plan: dict[str, dict] | None = None,
) -> str:
    """Construit un fragment OOXML <w:footnote> sans re-sérialiser tout footnotes.xml."""
    note_image_plan = note_image_plan or {}
    parts = [
        f'<w:footnote w:type="normal" w:id="{_xml_attr(footnote_id)}">',
        '<w:p>',
        '<w:pPr><w:pStyle w:val="Notedebasdepage"/></w:pPr>',
        '<w:r>',
        _note_rpr_xml(superscript=True),
        '<w:footnoteRef/>',
        '</w:r>',
        '<w:r>',
        _note_rpr_xml(),
        '<w:t xml:space="preserve">&#160;</w:t>',
        '</w:r>',
    ]

    doc_pr_id = 1000 * (footnote_id + 1)
    if note.inlines:
        for span in note.inlines:
            if span.kind == "note_call":
                continue
            if span.kind == "image":
                plan_entry = note_image_plan.get(span.attributes.get("occurrence_id", ""))
                if plan_entry is not None:
                    doc_pr_id += 1
                    parts.append(_note_image_drawing_xml(plan_entry["rid"], plan_entry["occurrence"], doc_pr_id))
                continue
            if not span.text:
                continue
            parts.append(
                _note_text_run_xml(
                    span.text,
                    bold=span.style.bold,
                    italic=span.style.italic,
                    small_caps=span.style.small_caps,
                    superscript=span.style.superscript,
                    highlight_key=span.attributes.get("highlight_color"),
                )
            )
    elif note.text:
        parts.append(_note_text_run_xml(note.text))

    parts.extend(['</w:p>', '</w:footnote>'])
    return "".join(parts)


_RELS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_IMAGE_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"


def _add_relationships_xml(rels_xml: bytes | None, entries: list[tuple[str, str]]) -> bytes:
    """Ajoute des <Relationship> (rid, target) de type image a un fichier .rels,
    en créant le fichier s'il n'existe pas encore."""
    if not entries:
        return rels_xml if rels_xml is not None else b""
    if rels_xml is None:
        rels_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\r\n'
            f'<Relationships xmlns="{_RELS_NS}"></Relationships>'
        ).encode("utf-8")
    fragment = "".join(
        f'<Relationship Id="{_xml_attr(rid)}" Type="{_IMAGE_REL_TYPE}" Target="{_xml_attr(target)}"/>'
        for rid, target in entries
    ).encode("utf-8")
    closing = b"</Relationships>"
    if closing not in rels_xml:
        raise ValueError("fichier .rels sans balise </Relationships> reconnaissable")
    return rels_xml.replace(closing, fragment + closing, 1)


def _ensure_content_type_defaults(content_types_xml: bytes, extensions: set[str]) -> bytes:
    missing = [ext for ext in extensions if f'Extension="{ext}"'.encode("utf-8") not in content_types_xml]
    if not missing:
        return content_types_xml
    ext_to_ct = {v: k for k, v in _CONTENT_TYPE_EXT.items()}
    fragment = "".join(
        f'<Default Extension="{_xml_attr(ext)}" ContentType="{_xml_attr(ext_to_ct.get(ext, "application/octet-stream"))}"/>'
        for ext in missing
    ).encode("utf-8")
    closing = b"</Types>"
    if closing not in content_types_xml:
        return content_types_xml
    return content_types_xml.replace(closing, fragment + closing, 1)


def _inject_footnotes(
    output_path: Path,
    notes: list[Note],
    note_id_map: dict[str, int],
    image_occurrences_by_id: dict[str, ImageOccurrence] | None = None,
    image_assets: dict[str, ImageAsset] | None = None,
) -> None:
    """
    Injecte les notes dans word/footnotes.xml sans re-sérialiser le fichier entier.

    On préserve ainsi les déclarations XML, les namespaces et les attributs mc:Ignorable
    du gabarit Métopes. C'est plus robuste que ElementTree.tostring(root), qui renomme
    les préfixes w14/w15/wp14 en ns1/ns2 et peut déclencher la réparation de Word.

    Les images incorporées dans les notes sont ajoutées ici (media, relations,
    content-types) car footnotes.xml est déjà écrit hors du modèle objet python-docx.
    """
    image_occurrences_by_id = image_occurrences_by_id or {}
    image_assets = image_assets or {}

    with zipfile.ZipFile(output_path, "r") as z:
        all_files = {name: z.read(name) for name in z.namelist()}

    fn_xml = all_files.get("word/footnotes.xml") or (_FOOTNOTES_XML_DECL + _FOOTNOTES_FALLBACK_ROOT)
    closing_tag = b"</w:footnotes>"
    if closing_tag not in fn_xml:
        raise ValueError("word/footnotes.xml ne contient pas de balise </w:footnotes> reconnaissable")

    existing_rels = all_files.get("word/_rels/footnotes.xml.rels")
    existing_rids = {
        m.decode("ascii") for m in re.findall(rb'Id="(rId\d+)"', existing_rels or b"")
    }
    note_image_plan = _build_note_image_plan(notes, image_occurrences_by_id, image_assets, existing_rids)

    fragments: list[str] = []
    for note in notes:
        fn_id = note_id_map.get(note.note_id)
        if fn_id is None:
            continue
        fragments.append(_build_footnote_xml(note, fn_id, note_image_plan))

    insertion = "".join(fragments).encode("utf-8")
    all_files["word/footnotes.xml"] = fn_xml.replace(closing_tag, insertion + closing_tag, 1)

    if note_image_plan:
        # Target dans footnotes.xml.rels est relatif a word/, donc "media/xxx.ext".
        rel_entries = [(entry["rid"], entry["filename"]) for entry in note_image_plan.values()]
        all_files["word/_rels/footnotes.xml.rels"] = _add_relationships_xml(existing_rels, rel_entries)

        seen_media: set[str] = set()
        for entry in note_image_plan.values():
            media_path = f"word/{entry['filename']}"
            if media_path in seen_media:
                continue
            seen_media.add(media_path)
            all_files[media_path] = entry["asset"].data

        extensions = {entry["filename"].rsplit(".", 1)[-1] for entry in note_image_plan.values()}
        if "[Content_Types].xml" in all_files:
            all_files["[Content_Types].xml"] = _ensure_content_type_defaults(
                all_files["[Content_Types].xml"], extensions
            )

    with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as z:
        for name, data in all_files.items():
            z.writestr(name, data)


# ── Application des polices PURH sur les styles du document ──────────────────

def _apply_purh_fonts(doc: DocxDoc) -> None:
    """Remplace les polices du gabarit Métopes par les substituts PURH.

    Corps → Garamond (empattement)   |   Titraille → Calibri (sans empattement)
    """
    for style in doc.styles:
        if style.name in _BODY_STYLE_NAMES:
            style.font.name = _FONT_BODY
        elif style.name in _HEAD_STYLE_NAMES:
            style.font.name = _FONT_HEAD


# ── Exporteur principal ───────────────────────────────────────────────────────

class DocxExporter:
    """
    Exporte un Document (modèle interne) vers un fichier .docx :
    - styles de paragraphes Métopes appliqués ;
    - corrections surlignées en couleur selon leur type ;
    - notes de bas de page reconstruites.
    """

    def __init__(self, template_path: Path | None = None) -> None:
        # Résolu paresseusement (pas au chargement du module) : une absence de
        # gabarit ne doit faire échouer qu'un export réel, pas tout import du module.
        self.template_path = resolve_default_template(explicit_path=template_path)

    def export(self, document: Document, output_path: Path) -> Path:
        """Génère le fichier .docx et retourne son chemin."""
        template_buf = _dotm_to_docx_bytes(self.template_path)
        doc = DocxDoc(template_buf)
        self._clear_body(doc)
        _apply_purh_fonts(doc)

        # Carte note_id → entier (IDs numériques Word des footnotes)
        note_id_map: dict[str, int] = {
            note.note_id: i for i, note in enumerate(document.notes, start=1)
        }
        image_occurrences_by_id = {occ.occurrence_id: occ for occ in document.image_occurrences}
        image_assets = document.image_assets

        for block in document.blocks:
            if block.attributes.get("page_break_before"):
                _add_page_break(doc)
            if block.block_type == "table":
                _add_raw_table(doc, block, image_assets)
            else:
                _add_paragraph(doc, block, note_id_map, image_occurrences_by_id, image_assets)

        if not document.blocks:
            doc.add_paragraph("")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        # Injection des notes de bas de page (et de leurs images) en post-traitement
        if document.notes:
            _inject_footnotes(
                output_path, document.notes, note_id_map,
                image_occurrences_by_id=image_occurrences_by_id,
                image_assets=image_assets,
            )

        return output_path

    @staticmethod
    def _clear_body(doc: DocxDoc) -> None:
        """Supprime le contenu du gabarit pour le remplacer par le document."""
        body = doc.element.body
        to_remove = [c for c in body if c.tag != qn("w:sectPr")]
        for el in to_remove:
            body.remove(el)
